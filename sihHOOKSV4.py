#!/usr/bin/env python3
import os
import subprocess
import argparse
from datetime import datetime
from openai import OpenAI
from docx import Document


# CONFIGURATION

OPENROUTER_API_KEY = "PUT-THE-KEY-HERE"  # <--- put your OpenRouter key here
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key=OPENROUTER_API_KEY
)

# Wordlists (adjust paths if needed)
SECLISTS = "/usr/share/seclists"
WORDLIST_DIRS = f"{SECLISTS}/Discovery/Web-Content/common.txt"
WORDLIST_VHOST = f"{SECLISTS}/Discovery/DNS/subdomains-top1million-5000.txt"
WORDLIST_XSS = f"{SECLISTS}/Fuzzing/XSS/xss-rsnake.txt"
WORDLIST_SQLI = f"{SECLISTS}/Fuzzing/SQLi/sql-injection.txt"


# SHELL RUNNER

def run_shell_command(command, output_file=None):
    print(f"[CMD] {command}")
    result = subprocess.run(command, shell=True, text=True,
                            capture_output=True)
    if result.stdout:
        print(result.stdout)
        if output_file:
            with open(output_file, "w") as f:
                f.write(result.stdout)
    if result.stderr:
        print(result.stderr)
    return result.stdout.strip()

# ==========================
# AI REPORT GENERATION
# ==========================
def generate_ai_report(aggregated_text, domain):
    print("[+] Sending data to OpenRouter for AI report generation...")
    resp = client.chat.completions.create(
        model="openai/gpt-4o-mini",
        messages=[
            {"role": "system", "content": "You are a cybersecurity expert. Write a clean, professional vulnerability report with clear sections: Recon, Fuzzing, Findings, Recommendations."},
            {"role": "user", "content": f"Generate a detailed vulnerability report for {domain}.

Findings:
{aggregated_text}"}
        ]
    )
    return resp.choices[0].message.content

# ==========================
# SAVE REPORT TO TXT
# ==========================
def save_report_docx(content, domain):
    report_dir = f"{domain}/report"
    os.makedirs(report_dir, exist_ok=True)
    report_file = os.path.join(report_dir, f"{domain}_report.txt")

    doc = Document()
    doc.add_heading(f"Security Assessment Report - {domain}", 0)
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("\n")
    doc.add_paragraph(content)
    doc.save(report_file)

    print(f"[+] Report saved: {report_file}")

# ==========================
# MAIN FUNCTION
# ==========================
def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("domain", help="Target domain (example.com)")
    parser.add_argument("--skip-nuclei", action="store_true", help="Skip nuclei scan")
    args = parser.parse_args()

    domain = args.domain
    base_dir = domain
    recon_dir = f"{base_dir}/recon"
    fuzzing_dir = f"{recon_dir}/fuzzing"
    wayback_dir = f"{recon_dir}/wayback"

    os.makedirs(recon_dir, exist_ok=True)
    os.makedirs(fuzzing_dir, exist_ok=True)
    os.makedirs(wayback_dir, exist_ok=True)

 
    # Step 1: Subfinder
 
    print("[+] Running Subfinder...")
    run_shell_command(f"subfinder -d {domain} -all -o {recon_dir}/subdomains.txt")

 
    # Step 2: Alive check (httpx)

    print("[+] Running HTTPX...")
    run_shell_command(f"httpx -l {recon_dir}/subdomains.txt -o {recon_dir}/alive.txt")


    # Step 3: Nuclei (optional)
 
    if not args.skip_nuclei:
        print("[+] Running Nuclei scan...")
        nuclei_templates = "/root/nuclei-templates/"
        run_shell_command(
            f"nuclei -l {recon_dir}/alive.txt -t {nuclei_templates} -o {recon_dir}/nuclei.txt"
        )
    else:
        print("[!] Skipping nuclei scan...")

 
    # Step 4: VHost fuzzing

    print("[+] Running VHost fuzzing...")
    run_shell_command(
        f'ffuf -w {/root/ffuf/SecLists/Discovery/DNS/subdomains-top1million-5000.txt} -u https://{domain}/ -H "Host: FUZZ.{domain}" -mc 200,403,500 -o {fuzzing_dir}/vhosts.txt -of md'
    )

  
    # Step 5: Directory fuzzing
  
    print("[+] Running FFUF directory fuzzing...")
    run_shell_command(
        f'ffuf -w {/root/ffuf/SecLists/Discovery/Web-Content/common.txt} -u https://{domain}/FUZZ -mc 200,403,500 -o {fuzzing_dir}/dirs.txt -of md'
    )


    # Step 6: Wayback data
  
    print("[+] Scraping Wayback data...")
    run_shell_command(f"waybackurls {domain} > {wayback_dir}/waybackdata.txt")

    print("[+] Extracting params and JS files...")
    with open(f"{wayback_dir}/waybackdata.txt") as f:
        lines = f.readlines()

    with open(f"{wayback_dir}/params.txt", "w") as p, \
         open(f"{wayback_dir}/jsfiles.txt", "w") as j:
        for line in lines:
            if "=" in line:
                p.write(line)
            if line.endswith(".js\n"):
                j.write(line)


    # Step 7: XSS fuzzing

    print("[+] Running XSS fuzzing...")
    run_shell_command(
        f'ffuf -w {wayback_dir}/params.txt:URL -w {/root/tools/Wordlists/fuzzdb/attack/xss/xss-rsnake.txt}:PAYLOAD -u "URLPAYLOAD" -mr PAYLOAD -mc 200,403,500 -o {fuzzing_dir}/xss.txt -of md'
    )


    # Step 8: SQLi fuzzing
   
    print("[+] Running SQLi fuzzing...")
    run_shell_command(
        f'ffuf -w {wayback_dir}/params.txt:URL -w {/root/ffuf/SecLists/Fuzzing/Databases/SQLi/Generic-BlindSQLi.fuzzdb.txt}:PAYLOAD -u "URLPAYLOAD" -mr PAYLOAD -mc 200,403,500 -o {fuzzing_dir}/sqli.txt -of md'
    )

  
    # Step 9: Aggregate results
  
    print("[+] Aggregating results for report...")
    aggregated_text = ""
    for path in [
        f"{recon_dir}/subdomains.txt",
        f"{recon_dir}/alive.txt",
        f"{recon_dir}/nuclei.txt",
        f"{fuzzing_dir}/vhosts.txt",
        f"{fuzzing_dir}/dirs.txt",
        f"{fuzzing_dir}/xss.txt",
        f"{fuzzing_dir}/sqli.txt",
    ]:
        if os.path.exists(path):
            with open(path, "r") as f:
                aggregated_text += f"\n\n==== {os.path.basename(path).upper()} ====
" + f.read()

   
    # Step 10: AI Report

    ai_report = generate_ai_report(aggregated_text, domain)
    save_report_docx(ai_report, domain)

    print("[+] Recon and report generation completed successfully!")


# ENTRYPOINT

if __name__ == "__main__":
    main()
